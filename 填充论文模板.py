from copy import deepcopy
from io import BytesIO
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
TEMPLATE_SOURCE = ROOT / "网站设计毕设论文.docx"
TARGET = ROOT / "网站设计毕设论文 - 副本.docx"
SOURCE = ROOT / "基于JavaScript的街舞信息共享平台设计与实现-毕业论文.docx"


def clear_paragraph(paragraph):
    """清空段落中的所有文字运行，保留段落本身结构和样式。"""
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_text(paragraph, text):
    """在保留模板段落样式的前提下替换整段文字。"""
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def set_cell_text(cell, text):
    """完整替换单元格内容，避免模板原文字残留。"""
    cell.text = text


def set_run_font(run, font_name=None, size=None, bold=None):
    """统一设置运行字体，并同步东亚字体配置，避免中文排版被 Word 回退。"""
    if font_name is not None:
        run.font.name = font_name
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(key), font_name)

    if size is not None:
        run.font.size = Pt(size)

    if bold is not None:
        run.bold = bold


def apply_exact_line_spacing(paragraph, size_pt):
    """按模板使用固定行高，减少标题和日期行在不同机器上的跳动。"""
    paragraph.paragraph_format.line_spacing = Pt(size_pt)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def copy_paragraph_style(source, target):
    """复制段落级样式，尽量复用模板与源论文已有格式。"""
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
    target.paragraph_format.line_spacing_rule = source.paragraph_format.line_spacing_rule
    target.paragraph_format.page_break_before = source.paragraph_format.page_break_before
    target.paragraph_format.keep_together = source.paragraph_format.keep_together
    target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
    target.paragraph_format.widow_control = source.paragraph_format.widow_control


def clone_run(source_run, target_paragraph):
    """复制运行级文本格式，使标题和正文更接近源论文效果。"""
    new_run = target_paragraph.add_run(source_run.text)
    new_run.bold = source_run.bold
    new_run.italic = source_run.italic
    new_run.underline = source_run.underline
    new_run.font.size = source_run.font.size
    new_run.font.name = source_run.font.name
    if source_run._r.rPr is not None:
        new_run._r.get_or_add_rPr().append(deepcopy(source_run._r.rPr))
    return new_run


def append_paragraph_from_source(document, source_paragraph):
    """把源论文段落追加到模板末尾。"""
    target = document.add_paragraph()
    copy_paragraph_style(source_paragraph, target)

    if source_paragraph.runs:
        for run in source_paragraph.runs:
            clone_run(run, target)
    else:
        target.add_run(source_paragraph.text)
    return target


def append_table_from_source(document, source_table):
    """复制源论文中的表格节点。"""
    document.element.body.insert(-1, deepcopy(source_table._element))


def add_standard_table(document, rows):
    """按标准 Word 表格重建正文表格，避免源论文中的异常表结构。"""
    if not rows or not rows[0]:
        return

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.style = "Normal"
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = None
                paragraph.paragraph_format.space_after = None
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT

    return table


def find_body_child_index_by_text(document, text):
    """按段落纯文本定位模板中的关键位置。"""
    for index, child in enumerate(document.element.body.iterchildren()):
        if child.tag == qn("w:p"):
            texts = [node.text for node in child.iter(qn("w:t")) if node.text]
            if "".join(texts).strip() == text:
                return index
    raise RuntimeError(f"未能在模板中定位到“{text}”。")


def remove_body_children_from_index(document, start_index):
    """删除模板中指定位置及其之后的正文节点，保留封面和声明。"""
    body = document.element.body
    children = list(body.iterchildren())
    sect_pr = None
    if children and children[-1].tag == qn("w:sectPr"):
        sect_pr = deepcopy(children[-1])

    for child in children[start_index:]:
        body.remove(child)

    if sect_pr is not None:
        body.append(sect_pr)


def remove_all_tables_after_first(document):
    """保留封面信息表，移除正文中旧项目遗留表格。"""
    for table in list(document.tables)[1:]:
        table._element.getparent().remove(table._element)


def set_cover_table(document):
    """按当前项目题目重写封面表格。"""
    table = document.tables[0]
    table.style = "Normal Table"
    table.autofit = False
    rows = [
        ("题    目", "基于JavaScript的街舞信息共享平台设计与实现"),
        ("专    业", ""),
        ("学生姓名", ""),
        ("班级学号", ""),
        ("指导教师", ""),
        ("指导单位", ""),
    ]
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.cell(index, 0), label)
        set_cell_text(table.cell(index, 1), value)

    # 封面表沿用模板的“左列标签、右列填写项”布局，右侧填写区统一居中，
    # 这样即使后续继续手动补信息，空位的视觉重心也会更接近模板。
    for row_index, row in enumerate(table.rows[: len(rows)]):
        left_paragraph = row.cells[0].paragraphs[0]
        right_paragraph = row.cells[1].paragraphs[0]

        left_paragraph.style = "Normal"
        left_paragraph.paragraph_format.first_line_indent = Pt(0)
        left_paragraph.alignment = None if row_index < 2 else WD_ALIGN_PARAGRAPH.CENTER

        right_paragraph.style = "Normal"
        right_paragraph.paragraph_format.first_line_indent = Pt(0)
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_basic_template_text(document):
    """替换模板封面、声明和日期中的旧项目信息。"""
    replacements = {
        "日期： 2024年1月8日至   2024年6月7日": "日期： 2026年3月20日至   2026年6月7日",
        "日期：2024年6月7日": "日期：2026年6月7日",
    }

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])
            if text.startswith("日期： 2024"):
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif text == "本人郑重声明：所提交的毕业设计（论文），是本人在导师指导下，独立进行研究工作所取得的成果。除文中已注明引用的内容外，本毕业设计（论文）不包含任何其他个人或集体已经发表或撰写过的作品成果。对本研究做出过重要贡献的个人和集体，均已在文中以明确方式标明并表示了谢意。":
            set_paragraph_text(
                paragraph,
                "本人郑重声明：所提交的毕业设计（论文），是在导师指导下独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本毕业设计（论文）不包含任何其他个人或集体已经发表或撰写过的作品成果。对本研究做出过重要贡献的个人和集体，均已在文中以明确方式标明并表示谢意。",
            )


def apply_cover_statement_format(document):
    """逐段修正封面页与声明页的关键格式，使副本更贴近模板版式。"""
    if len(document.paragraphs) < 15:
        return

    cover_date = document.paragraphs[6]
    cover_date.style = "Normal"
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_date.paragraph_format.first_line_indent = Pt(15)
    cover_date.paragraph_format.left_indent = None
    cover_date.paragraph_format.right_indent = None
    cover_date.paragraph_format.space_before = None
    cover_date.paragraph_format.space_after = None
    apply_exact_line_spacing(cover_date, 50)
    set_paragraph_text(cover_date, "日期： 2026年3月20日至   2026年6月7日")
    if cover_date.runs:
        set_run_font(cover_date.runs[0], font_name="楷体_GB2312", size=15)

    originality_title = document.paragraphs[8]
    originality_title.style = "Normal"
    originality_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    originality_title.paragraph_format.first_line_indent = Pt(0)
    originality_title.paragraph_format.left_indent = None
    originality_title.paragraph_format.right_indent = None
    originality_title.paragraph_format.space_before = None
    originality_title.paragraph_format.space_after = None
    set_paragraph_text(originality_title, "毕业设计（论文）原创性声明")
    if originality_title.runs:
        set_run_font(originality_title.runs[0], font_name="黑体", size=16)

    statement_text_map = {
        10: "本人郑重声明：所提交的毕业设计（论文），是在导师指导下独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本毕业设计（论文）不包含任何其他个人或集体已经发表或撰写过的作品成果。对本研究做出过重要贡献的个人和集体，均已在文中以明确方式标明并表示谢意。",
        13: "                                          论文作者签名：",
        14: "                                          日期：2026年6月7日",
    }
    for paragraph_index in range(10, 15):
        paragraph = document.paragraphs[paragraph_index]
        paragraph.style = "Normal"
        paragraph.alignment = None
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.right_indent = None
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if paragraph_index in statement_text_map:
            set_paragraph_text(paragraph, statement_text_map[paragraph_index])


def add_blank_paragraph(document, alignment=None):
    """补充用于控制章节节奏的空段，避免图表和正文粘连。"""
    paragraph = document.add_paragraph()
    paragraph.style = "Normal"
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.space_before = None
    paragraph.paragraph_format.space_after = None
    return paragraph


def add_table_caption(document, caption):
    """按模板中的表题样式补齐重建表格标题。"""
    paragraph = document.add_paragraph()
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.space_before = None
    paragraph.paragraph_format.space_after = None
    paragraph.add_run(caption)
    return paragraph


def add_cell_paragraph(cell, text, bold=False, font_name=None, size=None):
    """向单元格中追加一个居中段落，便于构造截图占位说明。"""
    paragraph = cell.add_paragraph()
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = None
    paragraph.paragraph_format.space_after = None
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold)
    return paragraph


def collect_blocks(source_document):
    """按照文档真实顺序收集段落和表格。"""
    paragraph_map = {paragraph._p: paragraph for paragraph in source_document.paragraphs}
    table_map = {table._element: table for table in source_document.tables}
    blocks = []
    for child in source_document.element.body.iterchildren():
        if child in paragraph_map:
            blocks.append(("paragraph", paragraph_map[child]))
        elif child in table_map:
            blocks.append(("table", table_map[child]))
    return blocks


def find_body_start_index(source_document):
    """定位源论文真正的第一章，而不是目录中的同名文字。"""
    for index, paragraph in enumerate(source_document.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if paragraph.text.strip() == "第一章 绪论" and style_name.startswith("Heading"):
            return index
    raise RuntimeError("未能在源论文中定位到“第一章 绪论”。")


def append_source_front(document, source_document, body_start_index):
    """写入摘要、英文摘要和目录标题位置，但目录本身改为可更新字段。"""
    source_paragraphs = source_document.paragraphs
    abstract_start = next(index for index, paragraph in enumerate(source_paragraphs) if paragraph.text.strip() == "摘  要")
    toc_started = False

    for paragraph in source_paragraphs[abstract_start:body_start_index]:
        text = paragraph.text.strip()
        if text == "目  录":
            toc_started = True
            target = append_paragraph_from_source(document, paragraph)
            append_toc_field(document)
            continue

        if toc_started:
            # 目录内容改由 Word 自动目录字段生成，因此跳过源论文中的静态目录文本。
            continue

        if text:
            append_paragraph_from_source(document, paragraph)


def append_toc_field(document):
    """插入 Word 自动目录字段，用户在 Word 中可直接更新目录。"""
    paragraph = document.add_paragraph()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "（右键目录后选择“更新域”即可生成目录）"
    placeholder_run.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run_begin = paragraph.add_run()
    run_begin._r.append(begin)

    run_instr = paragraph.add_run()
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    run_sep._r.append(separate)

    run_placeholder = paragraph.add_run()
    run_placeholder._r.append(placeholder_run)

    run_end = paragraph.add_run()
    run_end._r.append(end)


def add_centered_paragraph(document, text, bold=False, size=None):
    """追加一个居中段落，用于附录中的软件说明书封面。"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return paragraph


def add_code_heading(document, text):
    """追加附录中的小节标题。"""
    paragraph = document.add_paragraph()
    paragraph.style = "Heading 2"
    paragraph.add_run(text)
    return paragraph


def add_code_block(document, title, code):
    """把代码以多段正文形式写入 Word，便于保持模板的文本格式体系。"""
    add_code_heading(document, title)
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        paragraph.style = "Normal"
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        run = paragraph.add_run(line if line else " ")
        run.font.name = None
        run.font.size = None


def read_code_snippet(path, max_lines):
    """读取指定文件前若干行，作为附录代码展示素材。"""
    text = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    return "\n".join(text[:max_lines])


def append_appendix_manual(document):
    """将附录改为更贴近模板的软件说明书样式，并补充当前项目关键代码。"""
    document.add_page_break()
    appendix_title = add_centered_paragraph(document, "附录", bold=True, size=16)
    appendix_title.style = "Heading 1"
    appendix_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    appendix_title.paragraph_format.first_line_indent = Pt(0)

    school = add_centered_paragraph(document, "南京邮电大学", bold=True, size=14)
    apply_exact_line_spacing(school, 50)
    add_blank_paragraph(document, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    manual_title = add_centered_paragraph(document, "软 件 说 明 书", bold=True, size=18)
    manual_title.paragraph_format.first_line_indent = Pt(0)
    add_blank_paragraph(document)
    add_blank_paragraph(document)

    subject = add_centered_paragraph(document, "题  目：基于JavaScript的街舞信息共享平台设计与实现", bold=True, size=12)
    subject.paragraph_format.first_line_indent = Pt(0)
    add_blank_paragraph(document)
    document.add_paragraph("本附录用于补充正文中未完全展开的软件实现细节，主要展示街舞信息共享平台的关键前端页面代码、后端接口逻辑、数据模型设计以及演示数据初始化命令，便于说明系统的实际开发实现过程。")
    document.add_paragraph("附录中的代码均来源于当前项目工作目录中的真实文件，已经按照论文阅读场景做了适当截取，用于展示系统的主要实现方式与模块结构。")
    add_blank_paragraph(document)

    add_code_heading(document, "第一章 代码展示")

    code_map = [
        ("1.1 前端主路由与页面入口代码", ROOT / "frontend/src/App.jsx", 90),
        ("1.2 首页数据加载与平台介绍代码", ROOT / "frontend/src/pages/HomePage.jsx", 80),
        ("1.3 活动模块前端页面代码", ROOT / "frontend/src/pages/ActivitiesPage.jsx", 110),
        ("1.4 商城模块前端页面代码", ROOT / "frontend/src/pages/MallPage.jsx", 160),
        ("1.5 前端统一请求封装代码", ROOT / "frontend/src/api.js", 80),
        ("1.6 Django 项目总路由配置代码", ROOT / "config/urls.py", 80),
        ("1.7 用户模块核心视图代码", ROOT / "users/views.py", 160),
        ("1.8 活动模块核心视图代码", ROOT / "activities/views.py", 170),
        ("1.9 视频模块核心视图代码", ROOT / "videos/views.py", 170),
        ("1.10 社交模块核心视图代码", ROOT / "social/views.py", 140),
        ("1.11 商城模块核心视图代码", ROOT / "mall/views.py", 180),
        ("1.12 用户数据模型代码", ROOT / "users/models.py", 120),
        ("1.13 活动数据模型代码", ROOT / "activities/models.py", 120),
        ("1.14 视频数据模型代码", ROOT / "videos/models.py", 120),
        ("1.15 社交数据模型代码", ROOT / "social/models.py", 120),
        ("1.16 商城数据模型代码", ROOT / "mall/models.py", 140),
        ("1.17 演示数据初始化命令代码", ROOT / "users/management/commands/seed_demo_data.py", 220),
    ]

    for title, path, max_lines in code_map:
        add_code_block(document, title, read_code_snippet(path, max_lines))

    add_code_heading(document, "第二章 运行说明与接口摘要")
    document.add_paragraph("1. 系统运行方式：后端使用项目本地虚拟环境执行 migrate、seed_demo_data 和 runserver；前端进入 frontend 目录后执行 npm run dev。项目根目录还提供 start_all.ps1 一键启动脚本，可同时拉起前后端服务。")
    document.add_paragraph("2. 登录策略：系统支持游客浏览首页、活动、视频、社交和商城公开内容；当用户执行活动报名、商品购买、订单支付、评论互动和内容发布等需要个人身份信息的操作时，前端会提示先登录并跳转到登录页面。")
    document.add_paragraph("3. 主要接口分组：用户模块包括注册、登录、退出、个人信息维护、收藏和关注；活动模块包括活动列表、发布、报名、签到和收藏；视频模块包括视频列表、详情、点赞、收藏和评论；社交模块包括房间列表、房间详情和消息发送；商城模块包括商品、购物车、订单创建和模拟支付。")


def normalize_document_styles(document):
    """把现有论文内容重新套回模板里的标题、摘要、正文和代码格式。"""
    heading1_set = {
        "第一章 绪论",
        "第二章 设计方案",
        "第三章 制作过程分析",
        "第四章 总结与反思",
        "结束语",
        "致谢",
        "参考文献",
        "附录",
        "ABSTRACT",
        "第一章 代码展示",
        "第二章 运行说明与接口摘要",
    }
    center_title_set = {"摘  要", "目  录", "南京邮电大学", "软 件 说 明 书"}

    in_abstract = False
    in_english_abstract = False
    in_references = False
    in_appendix_manual = False

    def reset_to_body(p):
        p.style = "Normal"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.left_indent = None
        p.paragraph_format.space_before = None
        p.paragraph_format.space_after = None
        p.paragraph_format.page_break_before = None
        for run in p.runs:
            run.font.name = None
            run.font.size = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text == "摘  要":
            in_abstract = True
            in_english_abstract = False
            in_references = False
            in_appendix_manual = False
        elif text == "ABSTRACT":
            in_abstract = False
            in_english_abstract = True
            in_references = False
            in_appendix_manual = False
        elif text == "目  录":
            in_abstract = False
            in_english_abstract = False
            in_references = False
            in_appendix_manual = False
        elif text == "参考文献":
            in_abstract = False
            in_english_abstract = False
            in_references = True
            in_appendix_manual = False
        elif text == "附录":
            in_abstract = False
            in_english_abstract = False
            in_references = False
            in_appendix_manual = True
        elif text == "软 件 说 明 书":
            in_appendix_manual = True
        elif text in {"第一章 绪论", "第二章 设计方案", "第三章 制作过程分析", "第四章 总结与反思", "结束语", "致谢"}:
            in_abstract = False
            in_english_abstract = False
            in_references = False

        if text in {"南京邮电大学", "毕 业 设 计（论 文）", "题    目", "专    业", "学生姓名", "班级学号", "指导教师", "指导单位"} and not in_appendix_manual:
            continue

        if text in heading1_set:
            paragraph.style = "Heading 1"
            if text in {"附录", "参考文献", "ABSTRACT"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = None
            paragraph.paragraph_format.page_break_before = None
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            continue

        if text in center_title_set:
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            if paragraph.runs:
                for run in paragraph.runs:
                    if text == "摘  要":
                        run.font.size = Pt(16)
                    if text == "南京邮电大学" and in_appendix_manual:
                        run.bold = True
                    if text == "软 件 说 明 书":
                        run.bold = True
                        run.font.size = Pt(18)
            continue

        if text.startswith("关键词："):
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.page_break_before = None
            for run in paragraph.runs:
                run.bold = True
            continue

        if text.startswith("Key words"):
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.page_break_before = None
            continue

        if text.startswith("1.") or text.startswith("2.") or text.startswith("3.") or text.startswith("4."):
            dot_count = text.count(".")
            if dot_count == 1 and not text.endswith("。"):
                paragraph.style = "Heading 2"
                paragraph.alignment = None
                continue
            if dot_count >= 2 and not text.endswith("。"):
                paragraph.style = "Heading 3"
                paragraph.alignment = None
                continue

        if "论文插图示意" in text:
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = None
            paragraph.paragraph_format.space_after = None
            for run in paragraph.runs:
                run.italic = True
            continue

        if text.startswith("图"):
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = None
            paragraph.paragraph_format.space_after = None
            continue

        if text.startswith("表"):
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(21)
            paragraph.paragraph_format.space_before = None
            paragraph.paragraph_format.space_after = None
            continue

        if text.startswith("题  目："):
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                run.bold = True
            continue

        if in_abstract or in_english_abstract:
            reset_to_body(paragraph)
            paragraph.alignment = None
            continue

        if in_references and text.startswith("[") and "]" in text[:6]:
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.space_before = None
            paragraph.paragraph_format.space_after = None
            for run in paragraph.runs:
                run.font.name = None
                run.font.size = None
                run.bold = None
                run.italic = None
            continue

        if in_appendix_manual:
            reset_to_body(paragraph)
            paragraph.alignment = None
            continue

        code_like = (
            text.startswith("import ")
            or text.startswith("from ")
            or text.startswith("const ")
            or text.startswith("function ")
            or text.startswith("async ")
            or text.startswith("return ")
            or text.startswith("class ")
            or text.startswith("@")
            or text.startswith("def ")
            or text.startswith("if ")
            or text.startswith("elif ")
            or text.startswith("else")
            or text.startswith("try")
            or text.startswith("except")
            or text.startswith("with ")
            or text.startswith("urlpatterns")
            or text.startswith("export ")
            or text.startswith("<")
            or text.startswith("}")
            or text.startswith("{")
            or text.startswith("]")
            or text.startswith("[")
            or text.startswith(")")
            or text.startswith("(")
        )
        if code_like:
            reset_to_body(paragraph)
            paragraph.alignment = None
            continue

        reset_to_body(paragraph)


def normalize_safe_tables(document):
    """只对结构完整的表格做保守修正，尽量向模板表格靠拢。"""
    for table in document.tables:
        try:
            rows = table.rows
            cols = table.columns
            _ = len(rows)
            _ = len(cols)
        except Exception:
            continue

        try:
            first_text = table.cell(0, 0).text.strip()
        except Exception:
            continue

        # 封面表保留现有样式，仅把单元格首行缩进清零。
        if first_text == "题    目":
            try:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            except Exception:
                pass
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = "Normal"
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                        paragraph.paragraph_format.space_before = None
                        paragraph.paragraph_format.space_after = None
                if len(row.cells) >= 2:
                    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # 其余正文表尽量使用模板常见的 Table Grid，并让首行标题居中。
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = "Normal"
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = None
                    paragraph.paragraph_format.space_after = None
                    if row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def remove_invalid_tables(document):
    """删除无法被 python-docx 正常解析的异常表节点，只保留标准表格。"""
    for table in list(document.tables):
        try:
            _ = len(table.rows)
            _ = len(table.columns)
            _ = table.cell(0, 0).text
        except Exception:
            table._element.getparent().remove(table._element)


def add_diagram_figure(document, caption, title, rows):
    """插入程序开发常见风格的技术图表，如架构图、流程图、ER 图。"""
    add_blank_paragraph(document, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    table = document.add_table(rows=len(rows), cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row_index, row_text in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Pt(54 if row_index == 0 else 44)
        cell = row.cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        run = paragraph.add_run(row_text)
        set_run_font(run, font_name="宋体", size=10.5, bold=(row_index == 0))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Pt(0)
    run = caption_paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(10.5)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.first_line_indent = Pt(0)
    note.add_run(title).italic = True
    add_blank_paragraph(document)


def add_screenshot_placeholder(document, caption, page_name, capture_points):
    """插入系统截图占位框，方便用户后续手动截图替换。"""
    add_blank_paragraph(document, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.width = Inches(6.3)
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Inches(3.75)

    cell.text = ""
    header = cell.paragraphs[0]
    header.style = "Normal"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.first_line_indent = Pt(0)
    header.paragraph_format.space_before = None
    header.paragraph_format.space_after = None
    run = header.add_run("【系统截图手动插入区】")
    set_run_font(run, font_name="黑体", size=15, bold=True)

    add_cell_paragraph(cell, f"建议截图页面：{page_name}", bold=True, font_name="宋体", size=11)
    add_cell_paragraph(cell, "建议截图要求：完整页面、包含顶部/主内容区、分辨率清晰", font_name="宋体", size=10.5)

    add_cell_paragraph(cell, "截图内容建议：", bold=True, font_name="宋体", size=10.5)
    for item in capture_points:
        add_cell_paragraph(cell, f"· {item}", font_name="宋体", size=10.5)

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Pt(0)
    run = caption_paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(10.5)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.first_line_indent = Pt(0)
    note.add_run("此处需替换为系统实际运行截图，建议按说明手动截取后插入。").italic = True
    add_blank_paragraph(document)


TABLE_CAPTION_CONFIG = {
    "2.4 系统总体架构设计": "表2.1 系统总体架构设计表",
    "2.6 数据库表结构设计": "表2.2 数据库核心表结构设计表",
    "3.3 前后端接口联调": "表3.1 前后端主要接口联调表",
    "3.4 功能测试": "表3.2 系统功能测试结果表",
}


TECHNICAL_FIGURE_CONFIG = {}


SCREENSHOT_FIGURE_CONFIG = {
    "3.2.1 用户认证与个人中心实现": {
        "caption": "图3.1 用户登录与个人中心页面截图",
        "page_name": "登录页或个人中心页",
        "capture_points": [
            "登录/注册切换区域，能看出账号入口",
            "或个人中心页中的个人概况、统计卡片、资料编辑区域",
            "建议优先截个人中心页，信息量更丰富",
        ],
    },
    "3.2.2 活动模块实现": {
        "caption": "图3.2 活动模块页面截图",
        "page_name": "活动广场页面",
        "capture_points": [
            "活动列表卡片区域",
            "报名按钮、收藏按钮",
            "右侧或下方的发布活动表单",
        ],
    },
    "3.2.3 视频模块实现": {
        "caption": "图3.3 视频模块页面截图",
        "page_name": "视频作品墙页面",
        "capture_points": [
            "左侧作品列表",
            "右侧当前作品详情区域",
            "点赞、收藏、评论区域",
        ],
    },
    "3.2.4 社交聊天室实现": {
        "caption": "图3.4 社交聊天室页面截图",
        "page_name": "社交聊天室页面",
        "capture_points": [
            "左侧聊天室列表",
            "右侧聊天消息流",
            "底部消息输入与发送区域",
        ],
    },
    "3.2.5 商城模块实现": {
        "caption": "图3.5 商城模块页面截图",
        "page_name": "商城页面",
        "capture_points": [
            "商品卡片网格",
            "购物车或试选清单区域",
            "订单记录或下单相关区域",
        ],
    },
    "3.2.6 首页与整体交互实现": {
        "caption": "图3.6 首页整体交互页面截图",
        "page_name": "首页",
        "capture_points": [
            "首页平台简介区域",
            "聊天室分类、最新视频、商城商品概览区",
            "若有登录按钮或导航状态，尽量一起带上",
        ],
    },
}


REBUILT_TABLES = {
    "2.4 系统总体架构设计": [
        ["层次", "技术或模块", "主要职责"],
        ["表现层", "React、React Router、CSS3", "页面展示、路由跳转、表单输入、状态更新"],
        ["接口层", "fetch、JSON、Vite代理", "前端请求发送、错误提示、跨服务联调"],
        ["业务层", "Django Views、Session认证", "用户认证、活动报名、视频互动、商城订单等业务处理"],
        ["数据层", "Django ORM、SQLite", "实体建模、关系维护、数据持久化"],
    ],
    "2.6 数据库表结构设计": [
        ["表名", "核心字段", "说明"],
        ["users_user", "id、username、password、nickname、gender、phone、profile", "保存平台用户账号和扩展资料"],
        ["users_favorite", "id、user_id、target_type、target_id、created_at", "保存活动或视频收藏关系"],
        ["users_follow", "id、follower_id、following_id、created_at", "保存用户关注关系"],
        ["activities_activity", "id、organizer_id、title、activity_type、location、start_time、status", "保存活动基础信息"],
        ["activities_activityregistration", "id、activity_id、user_id、signup_time、checked_in", "保存活动报名和签到信息"],
        ["videos_video", "id、user_id、title、video_file、like_count、favorite_count、comment_count", "保存视频作品信息"],
        ["videos_videocomment", "id、video_id、user_id、content、created_at", "保存视频评论"],
        ["social_chatroom", "id、room_name、category、description、created_at", "保存聊天室分类"],
        ["social_chatmessage", "id、room_id、user_id、content、sent_at", "保存聊天室消息"],
        ["mall_product", "id、name、category、price、stock、status", "保存商城商品"],
        ["mall_cartitem", "id、user_id、product_id、quantity", "保存购物车项目"],
        ["mall_order", "id、user_id、total_amount、order_status、payment_status", "保存订单主表"],
        ["mall_orderitem", "id、order_id、product_id、quantity、unit_price", "保存订单明细"],
    ],
    "3.3 前后端接口联调": [
        ["模块", "主要接口", "说明"],
        ["用户", "/users/register/、/users/login/、/users/me/、/users/dashboard/", "注册登录、资料维护和个人中心统计"],
        ["活动", "/activities/list/、/activities/<id>/register/、/activities/<id>/checkin/", "活动列表、发布、报名和签到"],
        ["视频", "/videos/list/、/videos/<id>/like/、/videos/<id>/favorite/、/videos/<id>/comments/", "视频发布、点赞、收藏和评论"],
        ["社交", "/social/rooms/、/social/rooms/<id>/、/social/rooms/<id>/messages/", "聊天室列表、详情和消息发送"],
        ["商城", "/mall/products/、/mall/cart/、/mall/orders/create/、/mall/orders/<id>/pay/", "商品、购物车、订单和支付"],
    ],
    "3.4 功能测试": [
        ["测试编号", "测试内容", "预期结果", "测试结论"],
        ["T01", "访问系统首页", "首页显示平台简介、聊天室、视频和商品概览", "通过"],
        ["T02", "用户注册并自动登录", "注册成功后进入登录态，个人中心可访问", "通过"],
        ["T03", "使用错误密码登录", "系统提示用户名或密码错误", "通过"],
        ["T04", "未登录发布活动", "接口返回请先登录，页面显示错误提示", "通过"],
        ["T05", "发布合法活动", "活动创建成功并出现在活动列表", "通过"],
        ["T06", "报名同一活动两次", "第一次成功，第二次提示已报名", "通过"],
        ["T07", "活动签到", "报名用户签到成功并记录签到时间", "通过"],
        ["T08", "发布视频并评论", "视频出现在列表，评论显示在详情页", "通过"],
        ["T09", "收藏同一视频两次", "第一次成功，第二次提示已收藏", "通过"],
        ["T10", "进入聊天室发送消息", "消息保存并展示在房间消息列表中", "通过"],
        ["T11", "商品加入购物车并创建订单", "购物车清空，生成订单并扣减库存", "通过"],
        ["T12", "模拟支付订单", "订单状态更新为已支付", "通过"],
        ["T13", "修改个人资料", "资料保存成功，页面显示新昵称和简介", "通过"],
    ],
}


def append_source_body(document, source_document, body_start_index):
    """把第一章开始的正文与表格写入模板副本，并在关键章节后插入插图。"""
    blocks = collect_blocks(source_document)
    source_body_start_paragraph = source_document.paragraphs[body_start_index]
    body_started = False
    current_section_title = None
    rebuilt_done = set()

    for kind, block in blocks:
        if kind == "paragraph" and block._p is source_body_start_paragraph._p:
            body_started = True

        if not body_started:
            continue

        if kind == "paragraph":
            if block.text.strip() == "附录":
                # 附录不再沿用源论文里较简短的版本，而是改写成和模板更接近的软件说明书附录。
                append_appendix_manual(document)
                break

            appended = append_paragraph_from_source(document, block)
            title = block.text.strip()
            current_section_title = title
            if title in TECHNICAL_FIGURE_CONFIG:
                cfg = TECHNICAL_FIGURE_CONFIG[title]
                add_diagram_figure(document, cfg["caption"], cfg["title"], cfg["rows"])
            if title in SCREENSHOT_FIGURE_CONFIG:
                cfg = SCREENSHOT_FIGURE_CONFIG[title]
                add_screenshot_placeholder(document, cfg["caption"], cfg["page_name"], cfg["capture_points"])
            if title in REBUILT_TABLES and title not in rebuilt_done:
                add_table_caption(document, TABLE_CAPTION_CONFIG[title])
                add_standard_table(document, REBUILT_TABLES[title])
                add_blank_paragraph(document)
                rebuilt_done.add(title)
        elif kind == "table":
            if current_section_title in REBUILT_TABLES:
                # 当前章节的表格已由标准表重建，这里跳过源论文中原始异常表节点。
                continue
            append_table_from_source(document, block)


def main():
    """用原模板重建副本论文，并写入自动目录和示意图。"""
    if not TEMPLATE_SOURCE.exists():
        raise FileNotFoundError(f"未找到模板文件：{TEMPLATE_SOURCE}")
    if not SOURCE.exists():
        raise FileNotFoundError(f"未找到源论文内容文件：{SOURCE}")

    # 先用原模板重置副本，再在副本上直接写入，既保住模板版式，也满足用户要求。
    shutil.copyfile(TEMPLATE_SOURCE, TARGET)
    document = Document(TARGET)
    source_document = Document(SOURCE)

    set_cover_table(document)
    set_basic_template_text(document)

    template_abstract_index = find_body_child_index_by_text(document, "摘  要")
    source_body_start_index = find_body_start_index(source_document)

    remove_body_children_from_index(document, template_abstract_index)
    remove_all_tables_after_first(document)
    append_source_front(document, source_document, source_body_start_index)
    append_source_body(document, source_document, source_body_start_index)
    normalize_document_styles(document)
    apply_cover_statement_format(document)
    remove_invalid_tables(document)
    normalize_safe_tables(document)

    document.save(TARGET)
    print(f"已写入：{TARGET}")


if __name__ == "__main__":
    main()
